"""
Converter - Generate .md, .txt, or .docx reports from JSON output.

Usage:
    python converter.py reporte.json                     # -> reportes/reporte_YYYY-MM-DD_YYYY-MM-DD.md + .json
    python converter.py reporte.json --format md         # -> reportes/...md + .json
    python converter.py reporte.json --format docx       # -> reportes/...docx + .json
    python converter.py reporte.json --all               # -> reportes/...{.md,.txt,.docx,.json}
    python converter.py reporte.json -o my_report.md     # custom output (no reportes/ folder)
"""
import argparse
import shutil
import sys
from pathlib import Path

from docx import Document

from models import ReportOutput


# =============================================================================
# Output directory
# =============================================================================

REPORTES_DIR = Path(__file__).resolve().parent / "reportes"


# =============================================================================
# Formatters
# =============================================================================

def to_markdown(report: ReportOutput) -> str:
    """
    Convert a ReportOutput to a Markdown string.

    Args:
        report: Validated ReportOutput model

    Returns:
        Formatted Markdown string
    """
    lines: list[str] = []
    lines.append("# Reporte de Noticias - Sector CRE Mexico\n")
    lines.append(f"*Generado: {report.generated_at}*\n")
    lines.append(f"**Objetivo:** {report.objective}\n")
    lines.append(f"**Periodo:** {report.period_start} a {report.period_end}\n")
    lines.append("---\n")

    for section in report.digest.sections:
        lines.append(f"## {section.title}\n")
        lines.append(f"{section.article}\n")
        lines.append("**Fuentes:**\n")
        for source in section.sources:
            lines.append(f"- {source}")
        lines.append("\n---\n")

    return "\n".join(lines)


def to_plaintext(report: ReportOutput) -> str:
    """
    Convert a ReportOutput to a plain-text string.

    Args:
        report: Validated ReportOutput model

    Returns:
        Formatted plain-text string
    """
    sep = "=" * 60
    lines: list[str] = []
    lines.append(sep)
    lines.append("REPORTE DE NOTICIAS - SECTOR CRE MEXICO")
    lines.append(sep)
    lines.append(f"Generado:  {report.generated_at}")
    lines.append(f"Objetivo:  {report.objective}")
    lines.append(f"Periodo:   {report.period_start} a {report.period_end}")
    lines.append(sep)
    lines.append("")

    for i, section in enumerate(report.digest.sections, 1):
        lines.append(f"{i}. {section.title}")
        lines.append(f"   {section.article}")
        lines.append("   Fuentes:")
        for source in section.sources:
            lines.append(f"     - {source}")
        lines.append("")

    return "\n".join(lines)


def to_docx(report: ReportOutput) -> Document:
    """
    Convert a ReportOutput to a python-docx Document.

    Args:
        report: Validated ReportOutput model

    Returns:
        python-docx Document object (save with .save())
    """
    doc = Document()

    doc.add_heading("Reporte de Noticias - Sector CRE Mexico", level=0)

    doc.add_paragraph(f"Generado: {report.generated_at}")
    doc.add_paragraph(f"Objetivo: {report.objective}")
    doc.add_paragraph(f"Periodo: {report.period_start} a {report.period_end}")

    for section in report.digest.sections:
        doc.add_heading(section.title, level=1)
        doc.add_paragraph(section.article)

        doc.add_paragraph("Fuentes:", style="List Bullet")
        for source in section.sources:
            doc.add_paragraph(source, style="List Bullet 2")

    return doc


# Text-based formatters (return str)
TEXT_FORMATTERS = {
    "md": to_markdown,
    "txt": to_plaintext,
}

ALL_FORMATS = ["md", "txt", "docx"]


# =============================================================================
# Public API
# =============================================================================

def _report_basename(report: ReportOutput) -> str:
    """Derive a filename base from the report period, e.g. reporte_2026-02-07_2026-02-14."""
    return f"reporte_{report.period_start}_{report.period_end}"


def _save_format(report: ReportOutput, fmt: str, dest: Path) -> None:
    """Save a single format to dest path."""
    if fmt == "docx":
        doc = to_docx(report)
        doc.save(str(dest))
    else:
        content = TEXT_FORMATTERS[fmt](report)
        dest.write_text(content, encoding="utf-8")


def convert(
    json_path: str,
    fmt: str = "md",
    output_path: str | None = None,
    all_formats: bool = False,
) -> list[str]:
    """
    Read a JSON report and convert it to the requested format(s).

    By default, saves to reportes/ folder with period-based naming
    and copies the source JSON alongside it.

    Args:
        json_path: Path to the JSON report file
        fmt: Output format ('md', 'txt', or 'docx')
        output_path: Optional explicit output path (skips reportes/ logic)
        all_formats: If True, generate all formats (md, txt, docx)

    Returns:
        List of paths of written files
    """
    source = Path(json_path)
    if not source.exists():
        raise FileNotFoundError(f"No se encontró el archivo: {json_path}")

    report = ReportOutput.model_validate_json(source.read_text(encoding="utf-8"))
    created: list[str] = []

    # Explicit output path — simple mode (no reportes/ folder)
    if output_path and not all_formats:
        dest = Path(output_path)
        _save_format(report, fmt, dest)
        print(f"[OK] Convertido: {source} -> {dest}")
        return [str(dest)]

    # Default mode — save to reportes/ folder
    REPORTES_DIR.mkdir(exist_ok=True)
    basename = _report_basename(report)
    formats = ALL_FORMATS if all_formats else [fmt]

    for f in formats:
        if f not in TEXT_FORMATTERS and f != "docx":
            raise ValueError(f"Formato no soportado: {f}. Usa 'md', 'txt' o 'docx'.")
        dest = REPORTES_DIR / f"{basename}.{f}"
        _save_format(report, f, dest)
        print(f"[OK] Convertido: {source} -> {dest}")
        created.append(str(dest))

    # Copy JSON source to reportes/
    json_dest = REPORTES_DIR / f"{basename}.json"
    shutil.copy2(str(source), str(json_dest))
    print(f"[OK] Copiado:    {source} -> {json_dest}")
    created.append(str(json_dest))

    return created


# =============================================================================
# CLI Entry Point
# =============================================================================

def main() -> None:
    parser = argparse.ArgumentParser(
        description="Convierte reportes JSON del News Bot a Markdown, texto plano o Word."
    )
    parser.add_argument(
        "json_file",
        help="Ruta al archivo JSON generado por el agente",
    )
    parser.add_argument(
        "--format", "-f",
        choices=ALL_FORMATS,
        default="md",
        dest="fmt",
        help="Formato de salida (default: md)",
    )
    parser.add_argument(
        "--output", "-o",
        default=None,
        help="Ruta de salida personalizada (omite carpeta reportes/)",
    )
    parser.add_argument(
        "--all",
        action="store_true",
        dest="all_formats",
        help="Genera todos los formatos (md, txt, docx) en reportes/",
    )
    args = parser.parse_args()

    try:
        convert(args.json_file, args.fmt, args.output, args.all_formats)
    except (FileNotFoundError, ValueError) as e:
        print(f"[ERROR] {e}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
